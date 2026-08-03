import docx

doc = docx.Document()
doc.add_paragraph("Jane Doe")
doc.add_paragraph("jane.doe@example.com | +1 555-123-4567")

doc.add_paragraph("Summary")
doc.add_paragraph("Experienced backend engineer with a focus on scalable APIs.")

doc.add_paragraph("Skills")
doc.add_paragraph("Python, FastAPI, PostgreSQL, Docker, AWS, communication, leadership")

doc.add_paragraph("Experience")
doc.add_paragraph("Senior Backend Engineer at Acme Corp (2021-2026)")
doc.add_paragraph("Built and maintained microservices handling 1M+ requests/day.")

doc.add_paragraph("Education")
doc.add_paragraph("B.Tech in Computer Science, XYZ University, 2016-2020")

doc.save("_sample_resume.docx")
print("saved")
